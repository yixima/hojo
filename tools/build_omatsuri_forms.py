#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""おまつり歳時記プロポーザル 参加意向申出（9/7 17時必着）の提出用ファイルを作る。

原本 docs/omatsuri/04_yousikimatome_omaturi.docx から必要な様式だけを切り出し、
当社情報を差し込む。**様式の文言は1文字も書き換えない。**

────────────────────────────────────────────────────────────
2026-09-04 の作り直しで直した3つの欠陥（v1 で提出しかけたもの）
────────────────────────────────────────────────────────────
(1) **セクション設定を失っていた（最大の原因）**
    原本は4つのセクションからなる。第1号様式と誓約書はセクション1
    （w:docGrid type="lines"・余白 1134twips）に属するが、セクション区切りは
    「段落の w:pPr に埋め込まれた w:sectPr」として存在する。範囲外の段落を
    すべて削除したため、その段落ごと区切りが消え、**文書末尾のセクション4の設定
    （文字グリッド linesAndChars・余白 1418）が全体に適用された。**
    結果、文字が間延びし、1枚の様式が2ページに溢れた。
    → 切り出した範囲を支配する sectPr を探し、末尾の sectPr と差し替える。

(2) **差し込んだ値が右寄りの狭い枠で折り返していた**
    住所などの欄は左インデント 3990twips（70mm）で、残り幅は約100mm しかない。
    当社の住所は約121mm あり、必ず折り返す。
    → 欄のインデント／タブ位置を左へ寄せ、なお収まらなければ
      差し込んだ値の文字サイズだけを自動で落とす（fit_size）。

(3) **第1号様式の宛名に敬称が無かった**
    原本の第1号様式は「事務総長・代表理事　河村　正人」で止まっている。
    参考様式1・10-1 には「様」があり、協会側の様式で不統一になっている。
    → 第1号様式にも「　様」を補う。文言の書き換えではなく敬称の補記である。

**最大の教訓：v1 では docx を目で見ずに、python-docx の文字列抽出だけで
「検査を通過した」と報告していた。本スクリプトは LibreOffice で PDF に変換し、
ページ数・用紙・「折り返していないこと」を機械で検査する。**

使い方
  python3 tools/build_omatsuri_forms.py --date 2026-09-04 \
      --tanto-shozoku "…" --tanto-name "…" --tanto-tel "…" --tanto-mail "…"
"""
import argparse
import copy
import os
import re
import subprocess
import tempfile
import unicodedata

import docx
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

SRC = 'docs/omatsuri/04_yousikimatome_omaturi.docx'
OUT = 'docs/omatsuri/submit'

ADDRESS = '〒150-0001　東京都渋谷区神宮前6-18-10　海老名ビル4F'
NAME = '一般社団法人ジャパンプロモーション'
REPRESENTATIVE = '代表理事　生島　儀尊'

TWIP_MM = 25.4 / 1440.0
PT_MM = 25.4 / 72.0

RANGES = {
    'form1': (0, 37),       # 第1号様式（37 は改ページのみの段落なので含めない）
    'seiyaku': (179, 203),  # 参考様式1
    'himitsu': (378, 408),  # 参考様式10-1
}


def text_width_mm(s, pt):
    """全角＝1em、半角＝0.5em として幅を出す（安全側に多めに見積もる）。"""
    w = 0.0
    for ch in s:
        if ch == '\t':
            continue
        w += 1.0 if unicodedata.east_asian_width(ch) in 'WFA' else 0.5
    return w * pt * PT_MM


def fit_size(text, available_mm, base_pt, floor_pt=8.0, margin_mm=4.0):
    """available_mm に収まる最大の文字サイズ（0.5pt刻み）を返す。"""
    pt = base_pt
    while pt > floor_pt and text_width_mm(text, pt) > available_mm - margin_mm:
        pt -= 0.5
    return pt


def body_items(doc):
    return list(doc.element.body.iterchildren())


def governing_sectpr(items, start):
    """範囲 start 以降を支配する sectPr を返す。

    Word のセクションは「段落の w:pPr/w:sectPr」で区切られ、その段落までが
    1つのセクションになる。よって start 以降で最初に見つかる段落内 sectPr が
    この範囲を支配する。無ければ body 直下の sectPr。
    """
    for el in items[start:]:
        if el.tag == qn('w:p'):
            sect = el.find(qn('w:pPr') + '/' + qn('w:sectPr'))
            if sect is not None:
                return sect
        elif el.tag == qn('w:sectPr'):
            return el
    raise RuntimeError('sectPr が見つからない')


def keep_range(doc, start, end, sect_src):
    """[start, end) 以外を削除し、末尾の sectPr を sect_src の写しに差し替える。"""
    items = body_items(doc)
    body = doc.element.body
    keep_sect = copy.deepcopy(sect_src)
    for i, el in enumerate(items):
        if el.tag == qn('w:sectPr'):
            body.remove(el)
            continue
        if not (start <= i < end):
            el.getparent().remove(el)
    body.append(keep_sect)


def para(doc, items, idx):
    return Paragraph(items[idx], doc)


def set_indent(items, idx, left_tw, first_tw=0, strip_tabs=True):
    """左インデントを絶対値で設定し、行頭のタブを取り除く。"""
    p = items[idx]
    pPr = p.find(qn('w:pPr'))
    if pPr is None:
        pPr = p.makeelement(qn('w:pPr'), {})
        p.insert(0, pPr)
    old = pPr.find(qn('w:ind'))
    if old is not None:
        pPr.remove(old)
    ind = pPr.makeelement(qn('w:ind'), {})
    ind.set(qn('w:left'), str(int(left_tw)))
    ind.set(qn('w:leftChars'), '0')
    ind.set(qn('w:firstLine'), str(int(first_tw)))
    ind.set(qn('w:firstLineChars'), '0')
    pPr.append(ind)
    if strip_tabs:
        for r in p.findall(qn('w:r')):
            for t in r.findall(qn('w:tab')):
                r.remove(t)
            for t in r.findall(qn('w:t')):
                if t.text:
                    t.text = t.text.lstrip('\t')


def set_tab_pos(items, idx, pos_tw):
    pPr = items[idx].find(qn('w:pPr'))
    if pPr is None:
        return
    tabs = pPr.find(qn('w:tabs'))
    if tabs is None:
        return
    for tab in tabs.findall(qn('w:tab')):
        tab.set(qn('w:pos'), str(int(pos_tw)))


def run_size_pt(run, default=10.5):
    rPr = run._element.find(qn('w:rPr'))
    if rPr is not None:
        sz = rPr.find(qn('w:sz'))
        if sz is not None:
            return int(sz.get(qn('w:val'))) / 2.0
    return default


def para_size_pt(p, default=10.5):
    """段落の実効サイズ。run に sz が無くても段落の他の run から拾う。

    ここを間違えると幅の見積もりが小さく出て、「収まる」と誤判定する。
    実際 v1 の作り直しでは 14pt の欄を 10.5pt と誤認し、折り返しを見逃した。
    """
    sizes = []
    for r in p.runs:
        rPr = r._element.find(qn('w:rPr'))
        if rPr is not None:
            sz = rPr.find(qn('w:sz'))
            if sz is not None:
                sizes.append(int(sz.get(qn('w:val'))) / 2.0)
    return max(sizes) if sizes else default


def append_value(doc, items, idx, run_i, value, available_mm, margin_mm=4.0):
    """指定 run の直後に値を足し、収まらなければ**足した分だけ**小さくする。

    ラベル（原本の文字）のサイズは変えない。変えると様式の見た目が崩れるため。
    """
    p = para(doc, items, idx)
    # run_i に -1 を渡せるようにしつつ、あとで「挿入した run」を正しく指すため
    # ここで正の添字に直す（v1 では runs[-1+1]=runs[0] を縮めてしまい、
    # 差し込んだ値のサイズが一切変わらなかった）
    run_i = run_i % len(p.runs)
    label_run = p.runs[run_i]
    base = para_size_pt(p)
    label_w = sum(text_width_mm(r.text, run_size_pt(r, base)) for r in p.runs)
    pt = fit_size(value, available_mm - label_w, base, margin_mm=margin_mm)
    new = copy.deepcopy(label_run._element)
    for t in list(new):
        if t.tag in (qn('w:t'), qn('w:tab'), qn('w:br')):
            new.remove(t)
    t = new.makeelement(qn('w:t'), {})
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = value
    new.append(t)
    label_run._element.addnext(new)
    if pt != base:
        Paragraph(items[idx], doc).runs[run_i + 1].font.size = docx.shared.Pt(pt)
    return pt


def text_area_mm(sect):
    pgSz = sect.find(qn('w:pgSz'))
    pgMar = sect.find(qn('w:pgMar'))
    return (int(pgSz.get(qn('w:w'))) - int(pgMar.get(qn('w:left')))
            - int(pgMar.get(qn('w:right')))) * TWIP_MM


def jp_date(iso):
    y, m, d = iso.split('-')
    return f'{int(y)}年{int(m)}月{int(d)}日'


def build_form1(args, out, ind=2400, margin=4.0, drop_blanks=0):
    doc = docx.Document(SRC)
    items = body_items(doc)
    s, e = RANGES['form1']
    assert para(doc, items, s).text.strip() == '（第１号様式）', '原本の構成が変わっている'
    sect = governing_sectpr(items, s)
    width = text_area_mm(sect)

    p1 = para(doc, items, 1)
    p1.runs[0].text = '　' + jp_date(args.date)
    for r in p1.runs[1:]:
        r.text = ''

    # 宛名に敬称を補う（原本は「河村　正人」で止まっている。参考様式1・10-1 には様がある）
    p4 = para(doc, items, 4)
    p4.runs[-1].text = p4.runs[-1].text.rstrip() + '　様'

    IND = ind           # 原本 3990twips=70mm では当社の住所が必ず折り返す
    for idx in (7, 8, 9):
        set_indent(items, idx, IND)
    avail = width - IND * TWIP_MM
    append_value(doc, items, 7, -1, ADDRESS, avail, margin)
    append_value(doc, items, 8, -1, NAME, avail, margin)
    p9 = para(doc, items, 9)   # 末尾に「印」がある単一 run。印の手前に入れる
    p9.runs[0].text = f'代表者職氏名　{REPRESENTATIVE}　　　　　　印'

    set_indent(items, 31, 3000)
    for idx in (32, 33, 34, 35):
        set_indent(items, idx, 3200)
    avail2 = width - 3200 * TWIP_MM
    for idx, val in ((32, args.tanto_shozoku), (33, args.tanto_name),
                     (34, args.tanto_tel), (35, args.tanto_mail)):
        if val:
            append_value(doc, items, idx, -1, val, avail2, margin)

    keep_range(doc, s, e, sect)
    doc.save(out)
    return out


def build_seiyaku(args, out, ind=1680, margin=4.0, drop_blanks=0):
    """参考様式1（誓約書）。

    **原本のこの様式は2ページある。**2ページ目には
    「（共同企業体の場合は、代表企業が提出すること。）」の1行しかない。
    当社は単独応募でこの注記は当てはまらないが、様式の行は消さない。
    代わりに、**代表者職氏名の下にある空段落（余白）を必要な数だけ詰めて**
    注記を1ページ目に引き上げ、1枚に収める。文言は1文字も変えていない。
    """
    doc = docx.Document(SRC)
    items = body_items(doc)
    s, e = RANGES['seiyaku']
    assert para(doc, items, s).text.strip() == '（参考様式１）', '原本の構成が変わっている'
    sect = governing_sectpr(items, s)
    width = text_area_mm(sect)

    para(doc, items, 194).runs[0].text = '　　' + jp_date(args.date)

    # 所在地は ind、商号・代表者職氏名は行頭タブ3つ（840×3）で、いずれも
    # 2520twips から始まる。3行そろえて 1680twips へ寄せる
    IND = ind
    for idx in (197, 198, 199):
        set_indent(items, idx, IND)
    avail = width - IND * TWIP_MM
    append_value(doc, items, 197, -1, '　' + ADDRESS, avail, margin)
    append_value(doc, items, 198, -1, '　' + NAME, avail, margin)
    # runs[3] は「代表者職氏名」＋全角空白13個。後続 run に「○印」フィールドがあるので
    # 置換はせず、**余白用の空白だけ詰めてから**値を足す（詰めないと行が長くなり折り返す）
    p199 = para(doc, items, 199)
    p199.runs[3].text = '　代表者職氏名　'
    append_value(doc, items, 199, 3, REPRESENTATIVE + '　　', avail, margin)

    if drop_blanks:
        # 代表者職氏名（199）と注記（202）の間にある空段落を末尾から詰める
        removed = 0
        # 詰めてよいのは「文言を持たない空段落」だけ。誓約文・宛名・日付・
        # 記入欄・注記は1つも消さない
        for idx in (201, 200, 196, 195, 193, 192):
            if removed >= drop_blanks:
                break
            if not para(doc, items, idx).text.strip():
                items[idx].getparent().remove(items[idx])
                removed += 1
        items = [el for el in body_items(doc)]
        e -= removed

    keep_range(doc, s, e, sect)
    doc.save(out)
    return out


def build_himitsu(args, out, ind=2400, margin=4.0, drop_blanks=0):
    doc = docx.Document(SRC)
    items = body_items(doc)
    s, e = RANGES['himitsu']
    assert para(doc, items, s).text.strip() == '（参考様式１０-１）', '原本の構成が変わっている'
    sect = governing_sectpr(items, s)
    width = text_area_mm(sect)

    para(doc, items, 401).runs[0].text = '　　　　　　' + jp_date(args.date)

    TAB = ind           # 原本 3261twips=57mm では当社の住所が収まらない
    for idx in (403, 404, 405):
        set_tab_pos(items, idx, TAB)
    avail = width - TAB * TWIP_MM
    append_value(doc, items, 403, -1, '　' + ADDRESS, avail, margin)
    append_value(doc, items, 404, -1, '　' + NAME, avail, margin)
    # runs = [\t, '代表者職・氏', '名', 空白16個, 空白, '印'] ── 空白を詰めてから
    # 「名」の直後に値を入れる（詰めないと「印」が次ページへ押し出される）
    p405 = para(doc, items, 405)
    p405.runs[3].text = '　　　　'
    append_value(doc, items, 405, 2, '　' + REPRESENTATIVE, avail, margin)

    keep_range(doc, s, e, sect)
    doc.save(out)
    return out


def render_pdf(paths, outdir):
    """LibreOffice で PDF 化する。**提出物ではなく検査のためだけに使う**（L1 §5-3）。"""
    env = dict(os.environ)
    env['HOME'] = tempfile.mkdtemp(prefix='lo-')
    subprocess.run(['soffice', '--headless', '--norestore', '--convert-to', 'pdf',
                    '--outdir', outdir] + list(paths),
                   check=True, capture_output=True, env=env, timeout=600)
    return [os.path.join(outdir, os.path.splitext(os.path.basename(p))[0] + '.pdf')
            for p in paths]


def visual_lines(page, tol=3.0):
    """描画された文字を y 座標でまとめて「実際に見える1行」を復元する。

    pymupdf の line は、字間が広い箇所（様式のラベルは字送りが広い）で
    1行を複数に割ってしまう。**割れた断片を1行と誤認すると
    「折り返している」と誤判定する。**そこで y でまとめ直す。
    """
    spans = []
    for block in page.get_text('dict')['blocks']:
        for line in block.get('lines', []):
            for sp in line['spans']:
                if sp['text'].strip():
                    spans.append((round(sp['bbox'][3], 1), sp['bbox'][0], sp['text']))
    spans.sort()
    lines, cur, cur_y = [], [], None
    for y, x, t in spans:
        if cur_y is None or abs(y - cur_y) <= tol:
            cur.append((x, t)); cur_y = y if cur_y is None else cur_y
        else:
            lines.append(''.join(t for _, t in sorted(cur)))
            cur, cur_y = [(x, t)], y
    if cur:
        lines.append(''.join(t for _, t in sorted(cur)))
    return lines


def inspect(pdf, must_same_line):
    """ページ数・用紙・「同じ行に載っているべき語」を実測する。"""
    import pymupdf
    d = pymupdf.open(pdf)
    problems = []
    if d.page_count != 1:
        problems.append(f'{d.page_count} ページある（1枚に収まっていない）')
    w_mm = round(d[0].rect.width * PT_MM, 1)
    h_mm = round(d[0].rect.height * PT_MM, 1)
    if not (abs(w_mm - 210) < 1 and abs(h_mm - 297) < 1):
        problems.append(f'用紙が A4 ではない（{w_mm}×{h_mm}mm）')
    lines = []
    for page in d:
        lines += visual_lines(page)
    flat = [''.join(l.split()) for l in lines]
    for group in must_same_line:
        words = [''.join(g.split()) for g in group]
        ok = False
        for l in flat:
            pos, good = -1, True
            for w in words:                       # 同じ行に、この順で並んでいるか
                i = l.find(w, pos + 1)
                if i < 0:
                    good = False
                    break
                pos = i
            if good:
                ok = True
                break
        if not ok:
            problems.append('同じ行に載っていない: ' + ' → '.join(group))
    if '\ufffd' in ''.join(flat):
        problems.append('文字化け（U+FFFD）がある')
    return problems, lines


def check(path, groups, tmp):
    pdf = render_pdf([path], tmp)[0]
    problems, lines = inspect(pdf, groups)
    flat = ''.join(''.join(l.split()) for l in lines)
    strays = sorted(set(re.findall(
        r'（第[０-９]+号様式）|（参考様式[０-９]+(?:-[０-９]+)?）', flat)))
    if len(strays) != 1:
        problems.append('様式の混入または欠落: ' + str(strays))
    if not re.fullmatch(r'[A-Za-z0-9._-]+', os.path.basename(path)):
        problems.append('ファイル名が §7-11 に違反')
    return problems, strays, pdf


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('--date', default='2026-09-04')
    ap.add_argument('--tanto-shozoku', default='')
    ap.add_argument('--tanto-name', default='')
    ap.add_argument('--tanto-tel', default='')
    ap.add_argument('--tanto-mail', default='')
    ap.add_argument('--dump', action='store_true')
    args = ap.parse_args()

    if args.dump:
        d = docx.Document(SRC)
        for i, el in enumerate(body_items(d)):
            tag = el.tag.split('}')[1]
            print(i, tag, repr(Paragraph(el, d).text.strip()[:60] if tag == 'p' else ''))
        return

    os.makedirs(OUT, exist_ok=True)
    jobs = [
        (build_form1, 'omatsuri_01_sanka_ikou_moushide.docx', 2400, 1400, [
            ('参加意向申出書',), ('河村', '正人', '様'),
            ('住所', '〒150-0001', '海老名ビル4F'),
            ('商号又は名称', NAME),
            ('代表者職氏名', '生島', '儀尊', '印'),
            ('E-mail', 'ikushima@japanpromotion.org'),
        ]),
        (build_seiyaku, 'omatsuri_02_seiyakusho.docx', 1680, 900, [
            ('誓約書',), ('河村', '正人', '様'),
            ('所在地', '〒150-0001', '海老名ビル4F'),
            ('商号又は名称', NAME),
            ('代表者職氏名', '生島', '儀尊'),
        ]),
        (build_himitsu, 'omatsuri_03_himitsu_hoji_seiyakusho.docx', 2400, 1200, [
            ('河村', '正人', '様'),
            ('所在地', '〒150-0001', '海老名ビル4F'),
            ('商号又は名称', NAME),
            ('代表者職・氏名', '生島', '儀尊', '印'),
        ]),
    ]

    tmp = tempfile.mkdtemp(prefix='verify-')
    ng = 0
    print('検査は「LibreOffice で PDF に変換 → 1ページか・A4か・折り返していないか」を実測する。')
    for fn, name, ind0, ind_floor, groups in jobs:
        path = os.path.join(OUT, name)
        # 合格するまで、①字下げを詰める ②それでも駄目なら差し込んだ値だけ小さくする
        for step in range(14):
            ind = max(ind_floor, ind0 - 200 * step)
            margin = 4.0 + max(0, step - 4) * 4.0
            fn(args, path, ind=ind, margin=margin, drop_blanks=min(4, step // 2))
            problems, strays, pdf = check(path, groups, tmp)
            if not problems:
                print(f'OK  {path}   様式: {",".join(strays)}'
                      f'   （字下げ {ind}twips・余裕 {margin:.0f}mm・'
                      f'空段落を{min(4, step // 2)}個詰めて・{step + 1} 回目で適合）')
                break
        else:
            print(f'NG  {path}')
            for p in problems:
                print('      - ' + p)
            ng += len(problems)
    print(f'\n検査用 PDF: {tmp}')
    if ng:
        raise SystemExit(f'{ng} 件の不備がある。発行しない。')
    print('3件すべて検査を通過した（1ページ・A4・折り返しなし・文字化けなし）。')


if __name__ == '__main__':
    main()
