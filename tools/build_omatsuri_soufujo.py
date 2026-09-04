#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""おまつり歳時記 参加意向申出の送付状（郵送用）を作る。

送付状は協会の指定様式ではない（要領に定めがない）。郵送時に同封書類の
内訳を示すための添え状である。持参する場合は不要。

2026-09-04 の作り直しで直した点
  ・**用紙が US レター（215.9×279.4mm）だった。** python-docx の既定
    テンプレートがレターのため。**A4 を明示的に設定する。**
  ・同封書類の「1部」を全角空白で揃えていたため桁が崩れていた。
    **右揃えタブで揃える。**
  ・行が右端いっぱいまで伸びてはみ出しかけていた。余白を日本の
    ビジネス文書に合わせ、1行の長さを実測で検査する。
"""
import argparse
import os
import re
import subprocess
import tempfile

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Mm, Pt

OUT = 'docs/omatsuri/submit/omatsuri_00_soufujo.docx'
ADDRESS = '〒150-0001　東京都渋谷区神宮前6-18-10　海老名ビル4F'
NAME = '一般社団法人ジャパンプロモーション'
REPRESENTATIVE = '代表理事　生島　儀尊'
PT_MM = 25.4 / 72.0

ITEMS = [
    ('１．参加意向申出書（第１号様式）', '１部'),
    ('２．誓約書（参考様式１）', '１部'),
    ('３．業務実績を証明する書類（契約書の写し）', '１部'),
    ('４．業務説明資料提供申込書 兼 守秘義務誓約書（参考様式１０－１）', '１部'),
]


def build(date_iso, tanto_tel='', tanto_mail='', out=OUT):
    y, m, d = (int(v) for v in date_iso.split('-'))
    doc = docx.Document()

    sec = doc.sections[0]
    sec.page_width, sec.page_height = Mm(210), Mm(297)      # ← A4 を明示する
    sec.top_margin = sec.bottom_margin = Mm(25)
    sec.left_margin = sec.right_margin = Mm(25)
    body_mm = 210 - 25 - 25                                  # 本文幅 160mm

    st = doc.styles['Normal']
    st.font.name = 'ＭＳ 明朝'
    st.font.size = Pt(10.5)
    st.element.rPr.rFonts.set(qn('w:eastAsia'), 'ＭＳ 明朝')
    st.paragraph_format.line_spacing = 1.15
    st.paragraph_format.space_after = Pt(0)

    def p(text='', align=None, before=0, after=6):
        par = doc.add_paragraph(text)
        if align is not None:
            par.alignment = align
        par.paragraph_format.space_before = Pt(before)
        par.paragraph_format.space_after = Pt(after)
        return par

    R, C = WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.CENTER
    p(f'{y}年{m}月{d}日', R, after=18)
    p('公益社団法人２０２７年国際園芸博覧会協会', after=0)
    p('行催事部　行催事課　御中', after=18)
    p(ADDRESS, R, after=0)
    p(NAME, R, after=0)
    p(REPRESENTATIVE + '　　　　　　印', R, after=0)
    if tanto_tel or tanto_mail:
        p(f'電話　{tanto_tel}　E-mail　{tanto_mail}'.strip(), R, after=18)
    else:
        p('', after=18)
    p('２０２７年国際園芸博覧会　主催者催事「おまつり歳時記プロジェクト（仮）」', C, after=0)
    p('にかかる実施計画作成業務委託　参加意向申出書等の送付について', C, after=18)
    p('　拝啓　時下ますますご清栄のこととお慶び申し上げます。', after=4)
    p('　このたび、標記公募型プロポーザルにつきまして、下記のとおり参加意向申出書等を'
      '送付いたします。ご査収のほどよろしくお願い申し上げます。', after=4)
    p('　なお、業務説明資料の提供につきましても、参考様式１０－１を同封いたしましたので、'
      '併せてご高配を賜りますようお願い申し上げます。', after=4)
    p('敬具', R, after=14)
    p('記', C, after=12)
    # 「1部」は右揃えタブで揃える（全角空白で揃えると必ず桁がずれる）
    for label, count in ITEMS:
        par = doc.add_paragraph()
        par.paragraph_format.space_after = Pt(4)
        par.paragraph_format.left_indent = Mm(4)
        par.paragraph_format.tab_stops.add_tab_stop(Mm(body_mm - 4), WD_TAB_ALIGNMENT.RIGHT)
        par.add_run(f'{label}\t{count}')
    p('', after=8)
    p('以上', R, after=0)

    os.makedirs(os.path.dirname(out), exist_ok=True)
    doc.save(out)
    return out


def verify(path):
    env = dict(os.environ)
    env['HOME'] = tempfile.mkdtemp(prefix='lo-')
    tmp = tempfile.mkdtemp(prefix='verify-')
    subprocess.run(['soffice', '--headless', '--norestore', '--convert-to', 'pdf',
                    '--outdir', tmp, path], check=True, capture_output=True,
                   env=env, timeout=600)
    pdf = os.path.join(tmp, os.path.splitext(os.path.basename(path))[0] + '.pdf')
    import pymupdf
    d = pymupdf.open(pdf)
    problems = []
    if d.page_count != 1:
        problems.append(f'{d.page_count} ページある')
    w_mm = round(d[0].rect.width * PT_MM, 1)
    h_mm = round(d[0].rect.height * PT_MM, 1)
    if not (abs(w_mm - 210) < 1 and abs(h_mm - 297) < 1):
        problems.append(f'用紙が A4 ではない（{w_mm}×{h_mm}mm）')
    # 字間・タブで大きく空くと pymupdf は1行を複数に割る。**割れた断片を
    # 1行と誤認すると折り返し判定が壊れる**ので、y 座標でまとめ直す
    spans, right_edge = [], 0.0
    for block in d[0].get_text('dict')['blocks']:
        for line in block.get('lines', []):
            right_edge = max(right_edge, line['bbox'][2])
            for sp in line['spans']:
                if sp['text'].strip():
                    spans.append((round(sp['bbox'][3], 1), sp['bbox'][0], sp['text']))
    spans.sort()
    lines, cur, cur_y = [], [], None
    for yy, xx, tt in spans:
        if cur_y is None or abs(yy - cur_y) <= 3.0:
            cur.append((xx, tt))
            cur_y = yy if cur_y is None else cur_y
        else:
            lines.append(''.join(t for _, t in sorted(cur)))
            cur, cur_y = [(xx, tt)], yy
    if cur:
        lines.append(''.join(t for _, t in sorted(cur)))
    flat = [''.join(l.split()) for l in lines]
    for group in [('〒150-0001', '海老名ビル4F'), ('電話', 'E-mail'),
                  ('参加意向申出書（第１号様式）', '１部'),
                  ('業務説明資料提供申込書', '守秘義務誓約書', '（参考様式１０－１）', '１部')]:
        words = [''.join(g.split()) for g in group]
        ok = False
        for l in flat:
            pos, good = -1, True
            for w in words:
                i = l.find(w, pos + 1)
                if i < 0:
                    good = False
                    break
                pos = i
            if good:
                ok = True
                break
        if not ok:
            problems.append('同じ行に載っていない: ' + ' → '.join(group))
    if right_edge * PT_MM > 210 - 25 + 1:
        problems.append(f'本文が右余白をはみ出している（右端 {right_edge * PT_MM:.1f}mm）')
    if '�' in ''.join(flat):
        problems.append('文字化け（U+FFFD）がある')
    return problems, tmp


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('--date', default='2026-09-04')
    ap.add_argument('--tanto-shozoku', default='')
    ap.add_argument('--tanto-name', default='')
    ap.add_argument('--tanto-tel', default='')
    ap.add_argument('--tanto-mail', default='')
    a = ap.parse_args()
    out = build(a.date, a.tanto_tel, a.tanto_mail)
    problems, tmp = verify(out)
    if not re.fullmatch(r'[A-Za-z0-9._-]+', os.path.basename(out)):
        problems.append('ファイル名が §7-11 に違反')
    print(('OK  ' if not problems else 'NG  ') + out + f'   （検査用 PDF: {tmp}）')
    for p in problems:
        print('      - ' + p)
    if problems:
        raise SystemExit('不備があるため発行しない。')


if __name__ == '__main__':
    main()
